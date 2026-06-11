"""Render structured SCOTSMAN analysis results."""

from __future__ import annotations

from datetime import datetime
from html import escape
from io import BytesIO
from typing import Any


SCOTSMAN_ROWS = [
    ("S", "Solution", "Credible solution?"),
    ("C", "Competition", "Competitive position?"),
    ("O", "Originality", "Unique proposition?"),
    ("T", "Timescales", "Manageable bid timescale?"),
    ("S", "Size", "Right opportunity size?"),
    ("M", "Money", "Price within budget?"),
    ("A", "Authority", "Decision maker known?"),
    ("N", "Need", "Burning need?"),
]


def normalize_analysis(value: Any, request: dict[str, Any] | None = None) -> dict[str, Any]:
    """Return a predictable report dict even when the model returns partial data."""
    if not isinstance(value, dict):
        project_number = (request or {}).get("project_number") or (request or {}).get("project_id") or "unbekannt"
        return {
            "title": f"SCOTSMAN Bid-Qualifizierung Projekt #{project_number}",
            "decision": "ACTION REQUIRED",
            "total_score": 0,
            "decision_reason": str(value),
            "scorecard": [],
            "internal_evidence": [],
            "contacts": [],
            "next_steps": ["Analyse manuell pruefen, da kein strukturiertes Resultat erzeugt wurde."],
        }

    report = dict(value)
    report.setdefault("decision", "ACTION REQUIRED")
    report.setdefault("total_score", _score_total(report.get("scorecard")))
    report.setdefault("scorecard", [])
    report.setdefault("internal_evidence", [])
    report.setdefault("contacts", [])
    report.setdefault("next_steps", [])
    report.setdefault("metadata", _metadata_from_request(request))
    if not report.get("title"):
        project_number = (request or {}).get("project_number") or (request or {}).get("project_id") or "unbekannt"
        report["title"] = f"SCOTSMAN Bid-Qualifizierung Projekt #{project_number}"
    return report


def slack_summary_text(report: dict[str, Any], links: dict[str, str] | None = None) -> str:
    """Return the compact Slack message body."""
    decision = str(report.get("decision") or "ACTION REQUIRED").upper()
    total = report.get("total_score", 0)
    reason = str(report.get("decision_reason") or "").strip()
    lines = [
        f"*SCOTSMAN:* {decision} - {total}/32",
    ]
    if reason:
        lines.append(reason[:350])

    html_url = (links or {}).get("html_url")
    docx_url = (links or {}).get("docx_url")
    link_parts = []
    if html_url:
        link_parts.append(f"<{html_url}|HTML Report>")
    if docx_url:
        link_parts.append(f"<{docx_url}|DOCX Report>")
    if link_parts:
        lines.append(" | ".join(link_parts))
    return "\n".join(lines)


def render_html_report(report: dict[str, Any]) -> str:
    """Render a self-contained HTML SCOTSMAN report."""
    rows = _scorecard_rows(report)
    evidence = _list_items(report.get("internal_evidence"), ["Keine belastbare interne Evidenz gefunden."])
    contacts = _contact_rows(report.get("contacts"))
    steps = _list_items(report.get("next_steps"), ["Keine naechsten Schritte vorhanden."])
    title = escape(str(report.get("title") or "SCOTSMAN Bid-Qualifizierung"))
    decision = escape(str(report.get("decision") or "ACTION REQUIRED").upper())
    total = escape(str(report.get("total_score", 0)))
    reason = escape(str(report.get("decision_reason") or ""))
    decision_class = _decision_class(report.get("decision"))
    meta_items = _header_meta_items(report)
    return f"""<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Open+Sans:wght@400;600;700&display=swap" rel="stylesheet">
<style>
:root{{
  --m-ink:#221C35; --m-white:#FFFFFF; --m-night:#2B2049; --m-lavender:#F3EBFF;
  --m-purple:#422982; --m-mauve:#7B69A8; --m-lilac:#C59CFF; --m-lilac-soft:#D5BAFF;
  --m-coral:#FF675C; --m-rose:#E283AF;
  --ok-bg:var(--m-lavender); --ok-fg:var(--m-purple);
  --warn-bg:#FBEAF2; --warn-fg:#B4537F;
  --risk-bg:#FFEAE8; --risk-fg:#C2382E;
  --line:#E4DBF2; --zebra:#FAF7FF;
  --font:'Open Sans','Segoe UI',Arial,sans-serif;
}}
*{{box-sizing:border-box;margin:0;padding:0;}}
body{{font-family:var(--font);color:var(--m-ink);background:#EFEAF6;font-size:15px;line-height:1.55;
  -webkit-print-color-adjust:exact;print-color-adjust:exact;}}
.page{{max-width:880px;margin:32px auto;background:var(--m-white);box-shadow:0 2px 14px rgba(34,28,53,.10);}}
.header-band{{color:#fff;padding:40px 48px 32px;
  background:linear-gradient(315deg,#2E1A47 0%,#33245C 31%,#2B2049 39%,#1A1527 68%);}}
.header-band .eyebrow{{font-size:12px;font-weight:600;letter-spacing:.08em;text-transform:uppercase;color:var(--m-lilac);}}
.header-band h1{{font-size:30px;font-weight:700;margin:.25em 0 .2em;color:#fff;line-height:1.2;}}
.header-band .sub{{font-size:16px;color:rgba(255,255,255,.78);}}
.header-band .meta{{margin-top:18px;font-size:12px;color:rgba(255,255,255,.55);
  border-top:1px solid rgba(255,255,255,.18);padding-top:12px;display:flex;gap:18px;flex-wrap:wrap;}}
main{{padding:36px 48px 48px;}}
h2{{color:var(--m-purple);font-size:21px;font-weight:700;margin:36px 0 14px;
  border-bottom:2px solid var(--m-lilac-soft);padding-bottom:6px;}}
h2:first-of-type{{margin-top:0;}}
p{{margin-bottom:10px;}}
small,.caption{{font-size:12px;color:var(--m-mauve);}}
.decision{{display:flex;align-items:stretch;border:1px solid var(--risk-fg);border-radius:6px;overflow:hidden;margin-bottom:8px;}}
.decision .verdict{{background:var(--risk-bg);color:var(--risk-fg);font-size:30px;font-weight:700;
  display:flex;align-items:center;justify-content:center;padding:18px 28px;min-width:170px;}}
.decision .why{{padding:14px 20px;}}
.decision .why strong{{display:block;margin-bottom:2px;}}
.decision.ok{{border-color:var(--ok-fg);}}
.decision.ok .verdict{{background:var(--ok-bg);color:var(--ok-fg);}}
.decision.warn{{border-color:var(--warn-fg);}}
.decision.warn .verdict{{background:var(--warn-bg);color:var(--warn-fg);font-size:24px;}}
table{{border-collapse:collapse;width:100%;font-size:14px;table-layout:fixed;}}
th{{background:var(--m-lavender);color:var(--m-ink);font-weight:700;text-align:left;}}
th,td{{border:1px solid var(--line);padding:9px 12px;vertical-align:top;}}
tbody tr:nth-child(even) td{{background:var(--zebra);}}
td.center,th.center{{text-align:center;}}
.crit{{font-weight:700;}}
.crit small{{display:block;font-weight:400;font-style:normal;}}
.letter{{font-weight:700;color:var(--m-purple);text-align:center;background:var(--m-lavender)!important;}}
.chip{{display:inline-block;min-width:34px;text-align:center;font-weight:700;border-radius:4px;padding:2px 8px;font-size:13px;}}
.s-ok{{background:var(--ok-bg);color:var(--ok-fg);}}
.s-warn{{background:var(--warn-bg);color:var(--warn-fg);}}
.s-risk{{background:var(--risk-bg);color:var(--risk-fg);}}
tr.total td{{background:var(--m-lavender)!important;font-weight:700;}}
ol,ul{{margin:0 0 0 20px;}}
li{{margin-bottom:6px;}}
li::marker{{color:var(--m-purple);font-weight:700;}}
footer{{background:var(--m-night);color:rgba(255,255,255,.65);font-size:12px;
  padding:14px 48px;display:flex;justify-content:space-between;gap:16px;}}
footer .brand{{color:var(--m-lilac);font-weight:600;}}
@media print{{
  body{{background:#fff;}}
  .page{{box-shadow:none;margin:0;max-width:none;}}
  main{{padding:28px 0 32px;}}
  .header-band,footer{{padding-left:0;padding-right:0;}}
}}
@media (max-width:640px){{
  main{{padding:24px 20px 36px;}}
  .header-band{{padding:28px 20px 24px;}}
  .decision{{flex-direction:column;}}
  table{{table-layout:auto;}}
  footer{{padding:14px 20px;flex-direction:column;}}
}}
</style>
</head>
<body>
<div class="page">
  <div class="header-band">
    <div class="eyebrow">Bid-Qualifizierung &middot; SCOTSMAN</div>
    <h1>{title}</h1>
    <div class="sub">Mesoneer Opportunity Review</div>
    <div class="meta">{meta_items}</div>
  </div>

  <main>
    <div class="decision {decision_class}">
      <div class="verdict">{decision}</div>
      <div class="why">
        <strong>Total {total}/32</strong>
        {reason}
      </div>
    </div>

    <h2>SCOTSMAN-Bewertung</h2>
    <table>
      <thead>
        <tr>
          <th class="center" style="width:42px"></th>
          <th style="width:160px">Kriterium</th>
          <th class="center" style="width:80px">Score (0-4)</th>
          <th class="center" style="width:90px">Risiko</th>
          <th>Bewertung / Kommentar</th>
        </tr>
      </thead>
      <tbody>{rows}</tbody>
    </table>

    <h2>Interne Evidenz</h2>
    <ul>{evidence}</ul>

    <h2>Ansprechpartner</h2>
    <table>
      <thead>
        <tr><th style="width:200px">Person</th><th style="width:130px">Rolle</th><th>Zweck</th></tr>
      </thead>
      <tbody>{contacts}</tbody>
    </table>

    <h2>Naechste Schritte</h2>
    <ol>{steps}</ol>
  </main>

  <footer>
    <span class="brand">mesoneer</span>
    <span>SCOTSMAN Bid-Qualifizierung &middot; vertraulich</span>
  </footer>
</div>
</body>
</html>
"""


def render_docx_report(report: dict[str, Any]) -> bytes | None:
    """Render a DOCX report. Returns None when python-docx is unavailable."""
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        return None

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    doc.add_heading(str(report.get("title") or "SCOTSMAN Bid-Qualifizierung"), level=1)
    p = doc.add_paragraph()
    p.add_run(f"{str(report.get('decision') or 'ACTION REQUIRED').upper()}").bold = True
    p.add_run(f" | Total {report.get('total_score', 0)}/32")
    if report.get("decision_reason"):
        doc.add_paragraph(str(report["decision_reason"]))

    doc.add_heading("SCOTSMAN-Bewertung", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["#", "SCOTSMAN", "Beschreibung", "Bewertung 0-4", "Risiko", "Kommentar"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in _scorecard(report):
        cells = table.add_row().cells
        for idx, value in enumerate(
            [
                row.get("letter"),
                row.get("criterion"),
                row.get("description"),
                str(row.get("score", "")),
                row.get("risk"),
                row.get("comment"),
            ]
        ):
            cells[idx].text = str(value or "")

    _add_bullets(doc, "Interne Evidenz", report.get("internal_evidence"), "Keine belastbare interne Evidenz gefunden.")
    _add_bullets(doc, "Ansprechpartner", _contacts_as_strings(report.get("contacts")), "Keine belastbaren Ansprechpartner gefunden.")
    _add_bullets(doc, "Naechste Schritte", report.get("next_steps"), "Keine naechsten Schritte vorhanden.")

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _score_total(scorecard: Any) -> int:
    if not isinstance(scorecard, list):
        return 0
    total = 0
    for row in scorecard:
        if isinstance(row, dict):
            try:
                total += int(row.get("score") or 0)
            except (TypeError, ValueError):
                pass
    return total


def _metadata_from_request(request: dict[str, Any] | None) -> dict[str, Any]:
    if not isinstance(request, dict):
        return {}
    keys = ["offer_deadline", "qna_deadline", "contract_start"]
    return {key: request.get(key) for key in keys if request.get(key)}


def _header_meta_items(report: dict[str, Any]) -> str:
    metadata = report.get("metadata") if isinstance(report.get("metadata"), dict) else {}
    items = [
        ("Status", "Bewertet"),
        ("Q&A bis", _fmt_report_date(metadata.get("qna_deadline"))),
        ("Frist Einreichung", _fmt_report_date(metadata.get("offer_deadline"))),
        ("Start", _fmt_report_date(metadata.get("contract_start"))),
    ]
    visible = [(label, value) for label, value in items if value]
    if len(visible) == 1:
        visible.append(("Format", "mesoneer CI"))
    return "".join(f"<span>{escape(label)}: {escape(value)}</span>" for label, value in visible)


def _fmt_report_date(value: Any) -> str:
    if not value:
        return ""
    text = str(value)
    try:
        return datetime.fromisoformat(text.replace("Z", "+00:00")).strftime("%d.%m.%Y")
    except ValueError:
        try:
            return datetime.strptime(text, "%Y-%m-%d").strftime("%d.%m.%Y")
        except ValueError:
            return text


def _scorecard(report: dict[str, Any]) -> list[dict[str, Any]]:
    source = report.get("scorecard")
    rows_by_name = {}
    if isinstance(source, list):
        for row in source:
            if not isinstance(row, dict):
                continue
            key = str(row.get("criterion") or "").lower()
            if key:
                rows_by_name[key] = row

    rows = []
    for letter, criterion, description in SCOTSMAN_ROWS:
        row = dict(rows_by_name.get(criterion.lower(), {}))
        row.setdefault("letter", letter)
        row.setdefault("criterion", criterion)
        row.setdefault("description", description)
        row.setdefault("score", "")
        row.setdefault("risk", "")
        row.setdefault("comment", "")
        rows.append(row)
    return rows


def _scorecard_rows(report: dict[str, Any]) -> str:
    html_rows = []
    for row in _scorecard(report):
        score = row.get("score")
        risk = row.get("risk")
        html_rows.append(
            "<tr>"
            f'<td class="letter">{escape(str(row.get("letter") or ""))}</td>'
            f'<td class="crit">{escape(str(row.get("criterion") or ""))}'
            f'<small>{escape(str(row.get("description") or ""))}</small></td>'
            f'<td class="center"><span class="chip {_score_class(score)}">{escape(_display_value(score))}</span></td>'
            f'<td class="center"><span class="chip {_risk_class(risk)}">{escape(_display_value(risk))}</span></td>'
            f"<td>{escape(str(row.get('comment') or ''))}</td>"
            "</tr>"
        )
    html_rows.append(
        '<tr class="total">'
        "<td></td>"
        "<td>Total</td>"
        f'<td class="center">{escape(str(report.get("total_score", 0)))} / 32</td>'
        "<td></td>"
        f"<td>{escape(str(report.get('decision_reason') or ''))}</td>"
        "</tr>"
    )
    return "".join(html_rows)


def _list_items(values: Any, fallback: list[str]) -> str:
    if not isinstance(values, list) or not values:
        values = fallback
    return "".join(f"<li>{escape(_item_text(value))}</li>" for value in values[:6])


def _contact_items(values: Any) -> str:
    contacts = _contacts_as_strings(values)
    return _list_items(contacts, ["Keine belastbaren Ansprechpartner gefunden."])


def _contact_rows(values: Any) -> str:
    if not isinstance(values, list) or not values:
        return '<tr><td colspan="3">Keine belastbaren Ansprechpartner gefunden.</td></tr>'

    rows = []
    for value in values[:6]:
        if isinstance(value, dict):
            name = value.get("name") or value.get("person") or "Unbekannt"
            role = value.get("role") or ""
            reason = value.get("reason") or value.get("purpose") or ""
        else:
            name = str(value)
            role = ""
            reason = ""
        rows.append(
            "<tr>"
            f"<td>{escape(str(name or ''))}</td>"
            f"<td>{escape(str(role or ''))}</td>"
            f"<td>{escape(str(reason or ''))}</td>"
            "</tr>"
        )
    return "".join(rows)


def _contacts_as_strings(values: Any) -> list[str]:
    if not isinstance(values, list):
        return []
    result = []
    for value in values[:6]:
        if isinstance(value, dict):
            name = value.get("name") or value.get("person") or "Unbekannt"
            role = value.get("role") or ""
            reason = value.get("reason") or value.get("purpose") or ""
            result.append(" - ".join(part for part in [str(name), str(role), str(reason)] if part))
        else:
            result.append(str(value))
    return result


def _item_text(value: Any) -> str:
    if isinstance(value, dict):
        return " - ".join(str(part) for part in value.values() if part)
    return str(value)


def _display_value(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _decision_class(value: Any) -> str:
    decision = str(value or "").upper()
    if decision == "GO":
        return "ok"
    if decision == "ACTION REQUIRED":
        return "warn"
    return "risk"


def _score_class(value: Any) -> str:
    try:
        score = int(value)
    except (TypeError, ValueError):
        return "s-warn"
    if score >= 3:
        return "s-ok"
    if score >= 1:
        return "s-warn"
    return "s-risk"


def _risk_class(value: Any) -> str:
    risk = str(value or "").lower()
    if "sehr hoch" in risk or risk == "hoch":
        return "s-risk"
    if "mittel" in risk:
        return "s-warn"
    if "niedrig" in risk or "tief" in risk:
        return "s-ok"
    return "s-warn"


def _add_bullets(doc: Any, heading: str, values: Any, fallback: str) -> None:
    doc.add_heading(heading, level=2)
    if not isinstance(values, list) or not values:
        values = [fallback]
    for value in values[:6]:
        doc.add_paragraph(_item_text(value), style="List Bullet")
